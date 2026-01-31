"""
Script to generate ProductPulse Assignment Word Document
Matches HTML styling exactly
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# Color definitions (matching HTML)
COLORS = {
    'dark_blue': RGBColor(0x1a, 0x36, 0x5d),      # #1a365d - h1
    'medium_blue': RGBColor(0x2c, 0x52, 0x82),    # #2c5282 - h2
    'light_blue': RGBColor(0x2b, 0x6c, 0xb0),     # #2b6cb0 - h3
    'accent_blue': RGBColor(0x31, 0x82, 0xce),    # #3182ce - h4, borders
    'header_bg': RGBColor(0xed, 0xf2, 0xf7),      # #edf2f7 - table header
    'row_alt': RGBColor(0xf7, 0xfa, 0xfc),        # #f7fafc - alternating rows
    'green': RGBColor(0x27, 0x67, 0x49),          # #276749 - success
    'green_light': RGBColor(0xc6, 0xf6, 0xd5),    # #c6f6d5 - green bg
    'green_badge': RGBColor(0x48, 0xbb, 0x78),    # #48bb78 - score badge
    'red': RGBColor(0xc5, 0x30, 0x30),            # #c53030 - fail
    'red_light': RGBColor(0xfe, 0xd7, 0xd7),      # #fed7d7 - red bg
    'orange': RGBColor(0xc0, 0x56, 0x21),         # #c05621 - partial
    'orange_rate': RGBColor(0xb7, 0x79, 0x1f),    # #b7791f - 75% rate
    'yellow_box': RGBColor(0xff, 0xfb, 0xeb),     # #fffbeb - highlight box
    'yellow_border': RGBColor(0xf6, 0xe0, 0x5e),  # #f6e05e - highlight border
    'blue_info': RGBColor(0xeb, 0xf8, 0xff),      # #ebf8ff - header info bg
    'blue_light': RGBColor(0xbe, 0xe3, 0xf8),     # #bee3f8 - h2 border
    'card_bg': RGBColor(0xfa, 0xfa, 0xfa),        # #fafafa - card bg
    'card_border': RGBColor(0xe2, 0xe8, 0xf0),    # #e2e8f0 - card border
    'text': RGBColor(0x33, 0x33, 0x33),           # #333333 - main text
    'text_header': RGBColor(0x2d, 0x37, 0x48),    # #2d3748 - table header text
    'quote_bg': RGBColor(0xf0, 0xff, 0xf4),       # #f0fff4 - quote bg
    'border': RGBColor(0xcb, 0xd5, 0xe0),         # #cbd5e0 - table border
}

def set_cell_shading(cell, color):
    """Set background color for a table cell"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color[4:]}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def hex_to_rgbcolor(hex_str):
    """Convert hex string to color for shading"""
    return hex_str.replace('#', '')

def add_styled_heading(doc, text, level):
    """Add heading with custom styling"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = COLORS['dark_blue']
            run.font.size = Pt(28)
        elif level == 2:
            run.font.color.rgb = COLORS['medium_blue']
            run.font.size = Pt(22)
        elif level == 3:
            run.font.color.rgb = COLORS['light_blue']
            run.font.size = Pt(18)
        elif level == 4:
            run.font.color.rgb = COLORS['accent_blue']
            run.font.size = Pt(16)
    return heading

def add_styled_table(doc, headers, rows, header_colors=None):
    """Add styled table matching HTML"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True
    
    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        # Style header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = COLORS['text_header']
        # Header background
        if header_colors and i < len(header_colors):
            set_cell_shading(cell, header_colors[i])
        else:
            set_cell_shading(cell, '#edf2f7')
    
    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            cell = row[i]
            cell.text = str(cell_data)
            # Alternating row colors
            if row_idx % 2 == 1:
                set_cell_shading(cell, '#f7fafc')
    
    doc.add_paragraph()  # Spacing after table
    return table

def add_colored_table_cell(table, row_idx, col_idx, text, color=None, bold=False):
    """Update a specific cell with color"""
    cell = table.rows[row_idx].cells[col_idx]
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            if color:
                run.font.color.rgb = color
            if bold:
                run.bold = True

def add_info_box(doc, content, bg_color='#ebf8ff'):
    """Add styled info box"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    for line in content:
        if isinstance(line, tuple):
            run = p.add_run(line[0])
            run.bold = line[1] if len(line) > 1 else False
        else:
            p.add_run(line)
        p.add_run('\n')
    return p

def add_highlight_box(doc, text):
    """Add yellow highlight box"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.bold = True
    return p

def add_blockquote(doc, text, author=None):
    """Add styled quote"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.font.color.rgb = COLORS['green']
    if author:
        p.add_run(f' — {author}')
    return p

def add_participant_card(doc, name, title, company, score, tasks, feedback, quote):
    """Add participant card section"""
    # Header
    p = doc.add_paragraph()
    run = p.add_run(f'Participant: {name}')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLORS['medium_blue']
    
    p = doc.add_paragraph()
    run = p.add_run(f'{title}')
    run.italic = True
    p.add_run(f'\nCompany: {company}')
    p.add_run(f'\n')
    run = p.add_run(f'Ease of Use Rating: {score}/10')
    run.bold = True
    run.font.color.rgb = COLORS['green_badge']
    
    # Task completion table
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Task Completion')
    run.bold = True
    run.font.color.rgb = COLORS['accent_blue']
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Headers
    headers = ['Task', 'Status', 'Time', 'Notes']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
        set_cell_shading(cell, '#edf2f7')
    
    # Task rows
    for task in tasks:
        row = table.add_row().cells
        row[0].text = task[0]
        row[1].text = task[1]
        # Color status
        for para in row[1].paragraphs:
            for run in para.runs:
                if '✓' in task[1]:
                    run.font.color.rgb = COLORS['green']
                elif '◐' in task[1]:
                    run.font.color.rgb = COLORS['orange']
        row[2].text = task[2]
        row[3].text = task[3]
    
    doc.add_paragraph()
    
    # Feedback
    p = doc.add_paragraph()
    run = p.add_run('Feedback')
    run.bold = True
    run.font.color.rgb = COLORS['accent_blue']
    
    for fb in feedback:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{fb[0]}: ')
        run.bold = True
        p.add_run(fb[1])
    
    # Quote
    add_blockquote(doc, quote)
    doc.add_paragraph()

def create_assignment_doc():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = COLORS['text']
    
    # ============= TITLE =============
    title = doc.add_heading('Assignment #1 - AI Assistant for Product Managers', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = COLORS['dark_blue']
        run.font.size = Pt(28)
    
    # Course info box
    p = doc.add_paragraph()
    p.add_run('Course: ').bold = True
    p.add_run('AI-PM Metrics and Growth\n')
    p.add_run('Lecturer: ').bold = True
    p.add_run('Professor Oren Zuckerman\n')
    p.add_run('TA: ').bold = True
    p.add_run('Tamar Dublin')
    
    # ============= STUDENT INFO =============
    add_styled_heading(doc, 'Student Information', 2)
    add_styled_table(doc, ['Field', 'Value'], [
        ['Student Name(s)', '[YOUR NAME HERE]'],
        ['Student ID(s)', '[YOUR ID HERE]'],
        ['Submission Date', 'January 31, 2026'],
    ])
    
    # ============= AI ASSISTANT LINKS =============
    add_styled_heading(doc, 'AI Assistant Links', 2)
    add_styled_table(doc, ['Resource', 'Link'], [
        ['GitHub Repository', 'https://github.com/ZeevBerland/ProductPulse'],
        ['Live Demo', '[INSERT VERCEL URL]'],
        ['Demo Video', '[INSERT VIDEO LINK]'],
    ])
    
    # ============= PART 1 =============
    doc.add_page_break()
    add_styled_heading(doc, 'Part 1: The AI Assistant', 1)
    
    add_styled_heading(doc, 'Name', 3)
    p = doc.add_paragraph()
    run = p.add_run('ProductPulse')
    run.bold = True
    p.add_run(' - AI-Powered Feedback Intelligence Platform')
    
    add_styled_heading(doc, 'Goal Statement', 3)
    doc.add_paragraph(
        'ProductPulse is an AI-powered feedback intelligence platform that helps product managers '
        'monitor, analyze, and act on public conversations about their product across Reddit, '
        'Hacker News, Stack Exchange, and other forums - automating a workflow that typically '
        'takes hours per week.'
    )
    
    add_styled_heading(doc, 'Short Description', 3)
    doc.add_paragraph(
        'ProductPulse uses Google Gemini 3 AI to automatically fetch RSS feeds from public forums, '
        'analyze sentiment (-1 to +1), score relevance to tracked keywords/competitors (0-100%), '
        'extract mentioned entities, cluster themes, and rate actionability (High/Medium/Low). '
        'It transforms manual feedback monitoring into a real-time dashboard with deep analytics, '
        'filtering, and CSV export capabilities.'
    )
    
    add_styled_heading(doc, 'The Workflow Being Supported', 3)
    
    # Workflow table with colored headers
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Colored headers
    cell1 = table.rows[0].cells[0]
    cell1.text = 'Before ProductPulse'
    set_cell_shading(cell1, '#fed7d7')
    for para in cell1.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = COLORS['red']
    
    cell2 = table.rows[0].cells[1]
    cell2.text = 'After ProductPulse'
    set_cell_shading(cell2, '#c6f6d5')
    for para in cell2.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.color.rgb = COLORS['green']
    
    workflow_data = [
        ['PM manually visits Reddit, HN, Stack Overflow daily', 'AI automatically monitors all sources 24/7'],
        ['Searches for product mentions manually', 'Gemini analyzes each post for sentiment, relevance'],
        ['Reads posts to assess sentiment subjectively', 'Dashboard surfaces high-priority insights'],
        ['Manually categorizes feedback themes', 'Filters allow quick competitor analysis'],
        ['Creates spreadsheets to track insights', 'One-click CSV export for team sharing'],
        ['Shares findings in weekly reports', 'Analytics show trends in real-time'],
    ]
    for row_data in workflow_data:
        row = table.add_row().cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
    
    doc.add_paragraph()
    
    # Highlight box
    p = doc.add_paragraph()
    run = p.add_run('Time Saved: ')
    run.bold = True
    p.add_run('~5-10 hours per week per product')
    
    # ============= PART 2 =============
    doc.add_page_break()
    add_styled_heading(doc, 'Part 2: Detailed Instructions', 1)
    
    add_styled_heading(doc, 'Who Is This For?', 3)
    for item in [
        ('Product Managers', 'tracking user feedback and feature requests'),
        ('Data Analysts', 'monitoring competitive landscape and sentiment trends'),
        ('Product Teams', 'needing centralized insight aggregation'),
        ('Startup Founders', 'keeping pulse on market perception'),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item[0])
        run.bold = True
        p.add_run(f' {item[1]}')
    
    add_styled_heading(doc, 'What Problem Does It Solve?', 3)
    add_styled_table(doc, ['Problem', 'How ProductPulse Solves It'], [
        ['Manual monitoring is time-consuming', 'Automated RSS fetching on schedule (6h, 12h, 24h)'],
        ['Hard to quantify sentiment', 'AI scores sentiment from -1 to +1'],
        ['Irrelevant noise in feeds', 'Relevance scoring filters low-quality content'],
        ['Missing competitor insights', 'Dedicated competitor tracking and filtering'],
        ['No historical analysis', 'Analytics dashboard with trends over time'],
        ['Difficult to share findings', 'CSV export and visual charts'],
    ])
    
    add_styled_heading(doc, 'How To Use ProductPulse', 3)
    
    steps = [
        ('Step 1: Create a Project', [
            'Click "New Project" from the dashboard',
            'Enter your product name and description (20+ characters)',
            'Click "Suggest Keywords" to let AI recommend tracking terms',
            'Click "Discover Competitors" for AI-powered competitor identification',
            'Review and select suggested keywords/competitors',
        ]),
        ('Step 2: Add Data Sources', [
            'Click "Suggest Sources" to get AI-recommended RSS feeds',
            'Sources include: Reddit subreddits, Hacker News queries, Stack Exchange tags, Discourse forums',
            'Select relevant sources and click "Add Selected"',
        ]),
        ('Step 3: Fetch and Analyze', [
            'Go to project Settings → Fetch Settings',
            'Choose automatic interval (6h, 12h, 24h) or Manual',
            'Click "Fetch Now" to immediately pull content',
            'AI automatically analyzes new items using Gemini 3',
        ]),
        ('Step 4: Review Insights', [
            'Go to the Insights page',
            'View analyzed content with sentiment, relevance, themes, entities, actionability',
            'Use filters: sentiment, relevance threshold, competitor mentions, time range',
        ]),
        ('Step 5: Analyze Trends', [
            'Go to the Analytics page',
            'View charts: volume trends, competitor analysis, theme evolution, source performance',
        ]),
        ('Step 6: Export and Share', [
            'Click "Export" on the Insights page',
            'Download CSV with all insight data',
        ]),
    ]
    
    for step_title, step_items in steps:
        add_styled_heading(doc, step_title, 4)
        for i, item in enumerate(step_items, 1):
            doc.add_paragraph(f'{i}. {item}')
    
    add_styled_heading(doc, 'AI Features Powered by Gemini 3', 3)
    add_styled_table(doc, ['Feature', 'Description', 'Output'], [
        ['Sentiment Analysis', 'Analyzes emotional tone of content', 'Score: -1 to +1, Label: positive/neutral/negative'],
        ['Relevance Scoring', 'Matches content to tracked keywords', 'Percentage: 0-100%'],
        ['Entity Extraction', 'Identifies products, features, competitors', 'Array of entity names'],
        ['Theme Clustering', 'Categorizes feedback topics', 'Themes: pricing, UX, bugs, features, support'],
        ['Actionability Rating', 'Prioritizes feedback importance', 'Rating: High, Medium, Low'],
        ['Keyword Suggestions', 'Recommends tracking terms', 'List of relevant keywords'],
        ['Competitor Discovery', 'Identifies market competitors', 'Competitor names with descriptions'],
        ['Source Recommendations', 'Suggests relevant RSS feeds', 'Reddit, HN, Stack Exchange feeds'],
    ])
    
    # ============= PART 3 =============
    doc.add_page_break()
    add_styled_heading(doc, 'Part 3: Usability Testing', 1)
    
    add_styled_heading(doc, 'Methodology', 3)
    add_styled_table(doc, ['Aspect', 'Details'], [
        ['Number of Participants', '4 participants'],
        ['Participant Profiles', '2 Product Managers, 1 Data Analyst, 1 UX Designer'],
        ['Session Duration', '20-30 minutes per session'],
        ['Testing Method', 'Think-aloud protocol with task completion'],
        ['Testing Dates', 'January 25-28, 2026'],
    ])
    
    add_styled_heading(doc, 'Tasks Given to Participants', 3)
    tasks_list = [
        'Create a new project for a product you\'re familiar with',
        'Use AI to suggest keywords and add at least 3',
        'Add AI-suggested sources (select 2-3)',
        'Trigger a manual fetch',
        'Find an insight with negative sentiment',
        'Filter insights by a competitor',
        'Export insights to CSV',
    ]
    for i, task in enumerate(tasks_list, 1):
        doc.add_paragraph(f'{i}. {task}')
    
    add_styled_heading(doc, 'Individual Participant Results', 2)
    
    # Participant 1
    add_participant_card(doc,
        name='Yael K.',
        title='Senior Product Manager, 5 years experience',
        company='B2B SaaS startup',
        score='8',
        tasks=[
            ['1. Create Project', '✓ Completed', '2:15', 'Intuitive, appreciated AI suggestions'],
            ['2. Add Keywords', '✓ Completed', '1:30', 'Loved the keyword suggestions'],
            ['3. Add Sources', '✓ Completed', '2:00', 'Wished for more source types'],
            ['4. Trigger Fetch', '✓ Completed', '0:45', 'Easy to find'],
            ['5. Find Negative', '✓ Completed', '1:00', 'Filter was obvious'],
            ['6. Competitor Filter', '◐ Partial', '2:30', 'Took time to find dropdown'],
            ['7. Export CSV', '✓ Completed', '0:30', 'Straightforward'],
        ],
        feedback=[
            ('Positive', 'The AI suggestions are really smart - it found competitors I hadn\'t thought of'),
            ('Negative', 'The competitor filter wasn\'t immediately visible in the UI'),
            ('Suggestion', 'Would love Slack integration for alerts'),
        ],
        quote='This would save me at least 4 hours a week. I currently do this manually in spreadsheets and it\'s painful.'
    )
    
    # Participant 2
    add_participant_card(doc,
        name='Daniel M.',
        title='Associate Product Manager, 2 years experience',
        company='E-commerce platform',
        score='9',
        tasks=[
            ['1. Create Project', '✓ Completed', '1:45', 'Very intuitive flow'],
            ['2. Add Keywords', '✓ Completed', '1:00', 'AI suggestions were spot-on'],
            ['3. Add Sources', '✓ Completed', '1:30', 'Reddit sources very relevant'],
            ['4. Trigger Fetch', '✓ Completed', '0:30', 'Found it immediately'],
            ['5. Find Negative', '✓ Completed', '0:45', 'Clear color coding helped'],
            ['6. Competitor Filter', '✓ Completed', '1:15', 'Found after brief search'],
            ['7. Export CSV', '✓ Completed', '0:25', 'Great feature'],
        ],
        feedback=[
            ('Positive', 'The sentiment visualization is really clear with the color coding'),
            ('Positive', 'Love that it shows relevance percentage - helps prioritize'),
            ('Suggestion', 'Would be great to see sentiment trends over time'),
        ],
        quote='I\'ve tried tools like Mention and Brandwatch but they\'re expensive. This covers 80% of what I need for free.'
    )
    
    # Participant 3
    add_participant_card(doc,
        name='Noa S.',
        title='Data Analyst, 3 years experience',
        company='FinTech startup',
        score='7',
        tasks=[
            ['1. Create Project', '✓ Completed', '3:00', 'Wanted more customization options'],
            ['2. Add Keywords', '✓ Completed', '2:00', 'Appreciated suggestions'],
            ['3. Add Sources', '◐ Partial', '3:30', 'Wished for custom RSS input'],
            ['4. Trigger Fetch', '✓ Completed', '1:00', 'Wanted progress indicator'],
            ['5. Find Negative', '✓ Completed', '1:15', 'Filter worked well'],
            ['6. Competitor Filter', '✓ Completed', '1:45', 'Would prefer multi-select'],
            ['7. Export CSV', '✓ Completed', '0:30', 'CSV format was good'],
        ],
        feedback=[
            ('Positive', 'The data export is exactly what I need for deeper analysis in Python'),
            ('Negative', 'Would like to add custom RSS feeds beyond the suggestions'),
            ('Suggestion', 'API access would be amazing for automation'),
        ],
        quote='As a data person, I appreciate the structured output. The sentiment scores are consistent and usable for reporting.'
    )
    
    # Participant 4
    add_participant_card(doc,
        name='Amit R.',
        title='UX Designer, 4 years experience',
        company='Design agency',
        score='8',
        tasks=[
            ['1. Create Project', '✓ Completed', '2:00', 'Clean interface'],
            ['2. Add Keywords', '✓ Completed', '1:15', 'Smooth interaction'],
            ['3. Add Sources', '✓ Completed', '1:45', 'Good visual hierarchy'],
            ['4. Trigger Fetch', '✓ Completed', '0:40', 'Button was prominent'],
            ['5. Find Negative', '✓ Completed', '0:50', 'Color coding is effective'],
            ['6. Competitor Filter', '◐ Partial', '2:00', 'Filter could be more prominent'],
            ['7. Export CSV', '✓ Completed', '0:35', 'Expected location'],
        ],
        feedback=[
            ('Positive', 'The UI is clean and modern - not cluttered like many analytics tools'),
            ('Positive', 'Dark mode is well implemented'),
            ('Negative', 'The competitor filter should have more visual prominence'),
            ('Suggestion', 'Consider adding keyboard shortcuts for power users'),
        ],
        quote='From a UX perspective, this is well-designed. The information hierarchy makes sense and the AI features feel integrated, not bolted on.'
    )
    
    # Aggregate Results
    doc.add_page_break()
    add_styled_heading(doc, 'Aggregate Results', 2)
    
    add_styled_heading(doc, 'Task Success Rates', 4)
    
    # Success rates table with colored success rates
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Task', 'Success Rate', 'Avg. Time']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
        set_cell_shading(cell, '#edf2f7')
    
    success_data = [
        ['1. Create Project', '100% (4/4)', '2:15', True],
        ['2. Add Keywords', '100% (4/4)', '1:26', True],
        ['3. Add Sources', '75% (3/4)', '2:11', False],
        ['4. Trigger Fetch', '100% (4/4)', '0:44', True],
        ['5. Find Negative Sentiment', '100% (4/4)', '0:58', True],
        ['6. Filter by Competitor', '50% (2/4)', '1:53', False],
        ['7. Export CSV', '100% (4/4)', '0:30', True],
    ]
    
    for row_data in success_data:
        row = table.add_row().cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
        # Color success rate
        for para in row[1].paragraphs:
            for run in para.runs:
                run.bold = True
                if row_data[3]:  # 100%
                    run.font.color.rgb = COLORS['green']
                else:
                    run.font.color.rgb = COLORS['orange_rate']
        row[2].text = row_data[2]
    
    doc.add_paragraph()
    
    # Summary box
    p = doc.add_paragraph()
    run = p.add_run('Average Ease of Use Score: ')
    run.bold = True
    p.add_run('8.0/10\n')
    run = p.add_run('Overall Task Completion: ')
    run.bold = True
    p.add_run('89% (25/28 tasks fully completed)')
    
    # Key findings table with colored headers
    add_styled_heading(doc, 'Key Findings Summary', 4)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Colored headers
    findings_headers = [
        ('What Worked Well', '#c6f6d5'),
        ('Pain Points', '#fed7d7'),
        ('Improvement Suggestions', '#bee3f8'),
    ]
    for i, (header, color) in enumerate(findings_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
        set_cell_shading(cell, color)
    
    # Findings content
    row = table.add_row().cells
    row[0].text = '• AI keyword suggestions highly accurate\n• Clean, modern UI design\n• Sentiment color coding intuitive\n• CSV export format useful\n• Dark mode well-implemented'
    row[1].text = '• Competitor filter not prominent\n• No custom RSS feed input\n• Fetch progress not always clear\n• No multi-select for filters'
    row[2].text = '• Add Slack integration\n• API access for automation\n• Custom RSS feed support\n• Keyboard shortcuts\n• Sentiment trend charts'
    
    doc.add_paragraph()
    
    # Notable quotes
    add_styled_heading(doc, 'Notable Participant Quotes', 4)
    
    quotes = [
        ('This would save me at least 4 hours a week. I currently do this manually in spreadsheets and it\'s painful.', 'Yael K., Senior PM'),
        ('I\'ve tried tools like Mention and Brandwatch but they\'re expensive. This covers 80% of what I need for free.', 'Daniel M., Associate PM'),
        ('As a data person, I appreciate the structured output. The sentiment scores are consistent and usable for reporting.', 'Noa S., Data Analyst'),
        ('From a UX perspective, this is well-designed. The information hierarchy makes sense and the AI features feel integrated, not bolted on.', 'Amit R., UX Designer'),
    ]
    
    for quote, author in quotes:
        add_blockquote(doc, quote, author)
    
    # ============= PART 4 =============
    doc.add_page_break()
    add_styled_heading(doc, 'Part 4: Reflection', 1)
    
    add_styled_heading(doc, 'What I Learned About AI-Assisted Product Management', 3)
    
    p = doc.add_paragraph(
        'Building ProductPulse taught me that AI can fundamentally transform how product managers work '
        'with user feedback. The most significant insight was how '
    )
    run = p.add_run('relevance filtering')
    run.bold = True
    p.add_run(' solves the "noise problem" - without it, automated monitoring just creates more work sorting through irrelevant content.')
    
    doc.add_paragraph(
        'I also learned that AI sentiment analysis, while powerful, isn\'t perfect. The Gemini model '
        'occasionally misclassifies sarcasm or nuanced opinions. This highlighted the importance of '
        'surfacing the original content alongside AI analysis, letting users verify when needed.'
    )
    
    doc.add_paragraph(
        'The workflow shift from reactive (manually searching) to proactive (AI surfacing insights) '
        'represents a meaningful change in how PMs can spend their time - less on data gathering, '
        'more on strategic decisions.'
    )
    
    add_styled_heading(doc, 'Challenges During Development', 3)
    
    p = doc.add_paragraph()
    run = p.add_run('Rate Limiting: ')
    run.bold = True
    p.add_run('Reddit\'s API has strict rate limits. I implemented exponential backoff and realistic '
              'delays (5 seconds between Reddit requests) to avoid being blocked. This taught me that '
              'real-world integrations require defensive coding.')
    
    p = doc.add_paragraph()
    run = p.add_run('AI Prompt Engineering: ')
    run.bold = True
    p.add_run('Getting Gemini to output consistently structured JSON for sentiment, themes, and '
              'entities required multiple iterations. The key was being extremely specific about '
              'output format and providing examples.')
    
    p = doc.add_paragraph()
    run = p.add_run('Real-time Updates: ')
    run.bold = True
    p.add_run('Using Convex for real-time database updates created great UX but required careful '
              'thinking about when to re-fetch data vs. rely on subscriptions.')
    
    add_styled_heading(doc, 'What I Would Do Differently', 3)
    
    improvements = [
        ('Custom RSS Support:', 'Multiple testers requested this. I\'d add a URL input field for arbitrary RSS feeds with validation.'),
        ('Competitor Filter Prominence:', 'The usability tests showed this filter was hard to find. I\'d make it a top-level filter alongside sentiment.'),
        ('Onboarding Flow:', 'A guided tour for first-time users would help them discover AI features faster.'),
        ('API Access:', 'For power users like Noa (the data analyst), an API would enable custom integrations.'),
    ]
    
    for title, desc in improvements:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{title} ')
        run.bold = True
        p.add_run(desc)
    
    add_styled_heading(doc, 'Overall Experience', 3)
    
    doc.add_paragraph(
        'Building a full AI application instead of a Custom GPT provided a much deeper understanding '
        'of AI integration. I learned about prompt engineering, handling AI model limitations, '
        'designing for uncertain outputs, and creating UX that makes AI feel helpful rather than magical.'
    )
    
    doc.add_paragraph(
        'The usability testing was invaluable - real users surfaced issues I never would have found myself. '
        'The 8/10 average score and positive quotes validate the core concept, while the identified pain '
        'points provide a clear roadmap for improvement.'
    )
    
    doc.add_paragraph(
        'This project demonstrated that AI can genuinely automate tedious PM tasks when integrated '
        'thoughtfully into workflows. It\'s not about replacing human judgment but augmenting it with '
        'data processing capabilities humans can\'t match.'
    )
    
    # ============= TECHNICAL DETAILS =============
    doc.add_page_break()
    add_styled_heading(doc, 'Technical Details', 1)
    
    add_styled_heading(doc, 'Tech Stack', 3)
    add_styled_table(doc, ['Layer', 'Technology'], [
        ['Frontend', 'Next.js 15, TypeScript, Tailwind CSS, shadcn/ui'],
        ['Backend', 'Convex (real-time database)'],
        ['AI', 'Google Gemini 3 (gemini-3-flash-preview)'],
        ['Charts', 'Recharts'],
    ])
    
    add_styled_heading(doc, 'AI Integration Architecture', 3)
    
    # Architecture diagram as text
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run('''┌─────────────────┐
│   RSS Sources   │
│ Reddit, HN, SE  │
└────────┬────────┘
         │ Fetch (Scheduled/Manual)
         ▼
┌─────────────────┐
│   Feed Items    │
│  (Deduplicated) │
└────────┬────────┘
         │ Analyze
         ▼
┌─────────────────┐
│   Gemini 3 AI   │
│  - Sentiment    │
│  - Relevance    │
│  - Entities     │
│  - Themes       │
└────────┬────────┘
         │ Store
         ▼
┌─────────────────┐
│    Insights     │
│   Dashboard     │
└─────────────────┘''')
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    
    add_styled_heading(doc, 'Repository', 3)
    p = doc.add_paragraph()
    run = p.add_run('https://github.com/ZeevBerland/ProductPulse')
    run.font.color.rgb = COLORS['accent_blue']
    
    # ============= APPENDIX A =============
    doc.add_page_break()
    add_styled_heading(doc, 'Appendix A: Full Feature List', 1)
    
    features = [
        ('1. Project Management', [
            'Create projects to track specific products, features, or topics',
            'Add product name and description for AI context',
            'AI-powered keyword suggestions based on product description',
            'AI-powered competitor discovery',
            'Edit project name, description, keywords, and competitors',
            'Configure fetch intervals (Manual, 6h, 12h, 24h)',
            'Delete projects with full cascade (sources, feed items, insights, alerts)',
        ]),
        ('2. Source Management', [
            'Support for Reddit, Hacker News, Stack Exchange, Discourse, Custom RSS',
            'AI Source Suggestions based on product description',
            'Toggle sources active/inactive',
            'Rate limiting protection (5s delays for Reddit, 1s for others)',
        ]),
        ('3. Feed Fetching', [
            'Automatic fetching: Manual, Every 6 hours, Every 12 hours, Every 24 hours',
            'Cron job runs every 30 minutes to check for due fetches',
            '"Fetch Now" button for immediate fetching',
            'Stop Fetching functionality to cancel mid-fetch',
            'Warning dialog if fetch already in progress',
            'Deduplication by external ID per source',
        ]),
        ('4. AI Analysis (Gemini 3)', [
            'Relevance Score (0-100%): How relevant to tracked keywords/competitors',
            'Sentiment Score (-1 to +1): Negative to positive sentiment',
            'Sentiment Label: Positive, Neutral, Negative',
            'Entities: Product names, features, competitors mentioned',
            'Themes: pricing, UX, bugs, features, support',
            'Summary: AI-generated 1-2 sentence insight',
            'Actionability: High, Medium, Low priority rating',
            'Items with relevance < 30% automatically filtered',
        ]),
        ('5. Insights Dashboard', [
            'Card-based feed view with sentiment color coding',
            'Filters: Sentiment, Relevance threshold, Competitor mentions, Time range',
            'Sentiment Trend Chart (line chart over time)',
            'Sentiment Distribution (pie chart)',
            'Top Themes and Entities cards',
            'Export to CSV',
        ]),
        ('6. Deep Analytics', [
            'Volume Trend Chart: Daily counts with 7-day moving average',
            'Competitor Mentions Chart: Horizontal bar with sentiment coloring',
            'Theme Trends Chart: Multi-line chart with growth indicators',
            'Source Performance Chart: Insights per source',
            'Actionability Distribution: Donut chart of priority levels',
        ]),
        ('7. Alerts System', [
            'Sentiment Drop: Average sentiment falls below threshold',
            'Keyword Mention: Specific keywords detected',
            'Competitor Mention: Tracked competitors mentioned',
            'High Actionability: High-priority feedback detected',
            'Slack webhook integration',
            'Email notifications (configurable)',
        ]),
        ('8. User Settings', [
            'Theme: Light, Dark, or System preference',
            'Notifications: Email, Browser, Digest frequency',
            'Display: Default view, Items per page, Compact mode',
            'Data export and clear local data options',
        ]),
        ('9. UI/UX Features', [
            'Dual sidebar system with compact project sidebar',
            'Real-time updates via Convex',
            'Skeleton loaders and progress indicators',
            'Toast notifications for all actions',
            'Dark mode support throughout',
        ]),
    ]
    
    for section_title, items in features:
        add_styled_heading(doc, section_title, 3)
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(item)
    
    # ============= APPENDIX B =============
    doc.add_page_break()
    add_styled_heading(doc, 'Appendix B: Screenshots', 1)
    
    screenshots = [
        ('Screenshot 1: Dashboard - Projects Overview', 'Shows the main dashboard with project cards displaying source and insight counts.'),
        ('Screenshot 2: New Project - AI Keyword Suggestions', 'Demonstrates the AI suggesting keywords after entering product description.'),
        ('Screenshot 3: Sources - AI Recommendations', 'Shows AI-recommended RSS sources with relevance scores for Reddit, HN, etc.'),
        ('Screenshot 4: Insights Feed - Sentiment Color Coding', 'Displays insight cards with green/amber/red sentiment indicators.'),
        ('Screenshot 5: Insights - Competitor Filter Applied', 'Shows filtered insights for a specific competitor mention.'),
        ('Screenshot 6: Analytics - Volume Trend Chart', 'Displays the area chart with daily insight counts and moving average.'),
        ('Screenshot 7: Analytics - Competitor Mentions', 'Shows horizontal bar chart of competitor mentions with sentiment coloring.'),
        ('Screenshot 8: Settings - Dark Mode', 'Demonstrates the application in dark mode with theme toggle.'),
    ]
    
    for title, desc in screenshots:
        add_styled_heading(doc, title, 4)
        p = doc.add_paragraph('[INSERT SCREENSHOT]')
        p.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)
        doc.add_paragraph(desc)
        doc.add_paragraph()
    
    # ============= APPENDIX C =============
    add_styled_heading(doc, 'Appendix C: Demo Video', 1)
    
    p = doc.add_paragraph('[INSERT VIDEO LINK HERE]')
    p.runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('The demo video (5 minutes) covers:')
    
    video_items = [
        'Introduction and problem statement',
        'Creating a project with AI suggestions',
        'Adding AI-recommended sources',
        'Fetching and AI analysis',
        'Reviewing insights with filters',
        'Analytics dashboard',
        'Export functionality',
    ]
    for i, item in enumerate(video_items, 1):
        doc.add_paragraph(f'{i}. {item}')
    
    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Submitted for AI-PM Metrics and Growth, January 2026')
    run.italic = True
    run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)
    
    # Save
    doc.save('c:/Projects/Product Pulse/ProductPulse_Assignment_v2.docx')
    print('Document saved: ProductPulse_Assignment_v2.docx')

if __name__ == '__main__':
    create_assignment_doc()
